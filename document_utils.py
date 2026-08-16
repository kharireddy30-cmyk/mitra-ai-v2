import io
import docx

def extract_text_from_file(uploaded_file):
    """
    అప్‌లోడ్ చేసిన .docx లేదా .txt ఫైల్ నుండి టెక్స్ట్‌ను సంగ్రహిస్తుంది.
    """
    extracted = ""
    if uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        extracted = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif uploaded_file.name.endswith(".txt"):
        extracted = uploaded_file.read().decode("utf-8")
    return extracted

def create_docx_bytes(text):
    """
    టెక్స్ట్‌ను మైక్రోసాఫ్ట్ వర్డ్ (.docx) బైట్స్ రూపంలోకి మారుస్తుంది.
    """
    doc = docx.Document()
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def create_printable_pdf_html(text):
    """
    తెలుగు ఫాంట్ సపోర్ట్‌తో కూడిన ప్రింటబుల్ HTML/PDF టెంప్లేట్.
    """
    formatted_body = text.replace('\n', '<br>')
    return f"""<!DOCTYPE html>
<html lang="te">
<head>
    <meta charset="utf-8">
    <title>Spiritual Document</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Mandali&display=swap" rel="stylesheet">
    <style>
        body {{ 
            font-family: 'Mandali', Arial, sans-serif; 
            font-size: 18px; 
            line-height: 1.8; 
            padding: 40px; 
            color: #0f172a; 
        }}
        h2 {{ color: #1e3a8a; border-bottom: 2px solid #2563eb; padding-bottom: 8px; }}
        @media print {{ body {{ padding: 0; }} }}
    </style>
</head>
<body onload="window.print()">
    <h2>🕉️ ఆధ్యాత్మిక నోట్ / Spiritual Document</h2>
    <br>
    <div>{formatted_body}</div>
</body>
</html>"""

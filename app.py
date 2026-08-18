import streamlit as st
import io
import gc
from datetime import datetime
import docx
from streamlit_mic_recorder import speech_to_text
from services.audio_engine import transcribe_audio_file, synthesize_multilang_tts
from services.groq_polisher import polish_and_translate_script
from services.image_poster import generate_ai_poster_html

# ==========================================
# 1. పేజీ సెట్టింగ్స్ & స్టైల్స్
# ==========================================
st.set_page_config(
    page_title="BRAHMA AI Studio", 
    layout="wide", 
    page_icon="🕉️",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Mandali&display=swap');
    * { font-family: 'Mandali', 'Segoe UI', Tahoma, sans-serif; }
    .block-container { padding-top: 1rem; padding-bottom: 2rem; }
    div.stButton > button, div.stDownloadButton > button {
        font-weight: 600 !important;
        border-radius: 8px !important;
        padding: 6px 10px !important;
    }
    .diag-box {
        background-color: #0f172a;
        color: #38bdf8;
        border-radius: 6px;
        padding: 8px;
        font-family: 'Courier New', monospace;
        font-size: 11px;
        height: 90px;
        overflow-y: auto;
    }
    .diag-log { margin-bottom: 2px; }
</style>
""", unsafe_allow_html=True)

st.subheader("🕉️ BRAHMA AI : Studio (Multilingual Voice & Poster)")

# సెషన్ స్టేట్స్
if "main_text" not in st.session_state:
    st.session_state.main_text = ""
if "translated_text" not in st.session_state:
    st.session_state.translated_text = ""
if "audio_bytes_data" not in st.session_state:
    st.session_state.audio_bytes_data = None
if "poster_html_data" not in st.session_state:
    st.session_state.poster_html_data = None
if "last_mic_text" not in st.session_state:
    st.session_state.last_mic_text = ""
if "diag_logs" not in st.session_state:
    st.session_state.diag_logs = [
        {"time": datetime.now().strftime("%H:%M:%S"), "msg": "System Ready. Accurate Multilingual Translation Online.", "color": "#38bdf8"}
    ]

def add_log(msg, color="#38bdf8"):
    t_str = datetime.now().strftime("%H:%M:%S")
    st.session_state.diag_logs.append({"time": t_str, "msg": msg, "color": color})

def extract_text_from_file(uploaded_file):
    extracted = ""
    if uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        extracted = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif uploaded_file.name.endswith(".txt"):
        extracted = uploaded_file.read().decode("utf-8")
    return extracted

def create_docx_bytes(text):
    doc = docx.Document()
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def create_printable_pdf_html(text):
    formatted_body = text.replace('\n', '<br>')
    return f"""<!DOCTYPE html><html lang="te"><head><meta charset="utf-8"><title>Speech Script</title>
    <style>body {{ font-family: Arial, sans-serif; font-size: 17px; line-height: 1.8; padding: 25px; color: #000; }}</style></head>
    <body onload="window.print()"><div>{formatted_body}</div></body></html>"""


# ==========================================
# 2. AI స్పీచ్ స్టైల్ & పోస్టర్ సెట్టింగ్స్
# ==========================================
with st.expander("⚙️ AI CONTROLS (స్పీచ్ స్టైల్, థీమ్స్ & పోస్టర్ లేఅవుట్)", expanded=False):
    col_style, col_pause = st.columns(2)
    with col_style:
        selected_style = st.selectbox("🎭 స్పీచ్ స్టైల్:", options=["📢 పబ్లిక్ అనౌన్స్‌మెంట్ (Public Notice)", "🧘 ఆధ్యాత్మికం (Spiritual & Calm)", "📰 న్యూస్ రీడర్ (News Bulletin)", "🗣️ సంభాషణ / కబుర్లు (Conversational)"])
    with col_pause:
        selected_pause = st.selectbox("⏱️ శ్వాస విరామాలు:", options=["మధ్యస్థం (Normal Pauses)", "ఎక్కువ (Deep Breathing / Heavy Pauses)", "స్వల్పం (Fast / Light Pauses)"])

    col_th, col_stk = st.columns(2)
    with col_th:
        poster_theme = st.selectbox("🎨 పోస్టర్ కలర్ థీమ్:", options=["ఆధ్యాత్మికం (Golden Divine)", "రక్తదానం / సేవా కార్యక్రమం (Red & White)", "ప్రకృతి / పచ్చదనం (Nature Green)", "రాయల్ బ్లూ (Corporate & Formal)"])
    with col_stk:
        sticker_choice = st.selectbox(
            "🏷️ AI స్టిక్కర్స్ / బ్యాడ్జ్ ఎంపిక:",
            options=["🪄 AI మ్యాజిక్ (Auto Select)", "🕉️ ఓం (Divine Om)", "🪷 పద్మం (Sacred Lotus)", "🩸 రక్తదానం (Blood Drop)", "🕊️ శాంతి కపోతం (Peace Dove)", "🌟 గోల్డెన్ స్టార్ (Golden Star)", "📜 రాయల్ సీల్ (Royal Seal)", "❤️ సేవా హస్తం (Loving Care)"]
        )

    col_mode, col_align, col_fsize = st.columns(3)
    with col_mode:
        content_mode = st.selectbox("📝 కంటెంట్ మోడ్:", options=["📜 పూర్తి మ్యాటర్ (Full Exact Text)", "🤖 AI సారాంశం (Summary Points)"])
    with col_align:
        text_align = st.selectbox("📐 టెక్స్ట్ అమరిక (Alignment):", options=["ఎడమ వైపు (Left)", "మధ్యలో (Center)", "సమానంగా (Justify)"])
    with col_fsize:
        font_size_choice = st.selectbox("🔤 అక్షరాల సైజు (Font Size):", options=["మధ్యస్థం (Medium - 18px)", "చిన్నది (Small - 15px)", "పెద్దది (Large - 22px)", "చాలా పెద్దది (X-Large - 26px)"])

    custom_sticker_file = st.file_uploader("🖼️ కస్టమ్ లోగో/స్టిక్కర్ అప్‌లోడ్:", type=["png", "jpg", "jpeg", "webp"], key="cust_sticker_up")


# ==========================================
# 3. నాలుగు ఇన్‌పుట్ విభాగాలు (DOC | AUDIO | MIC | PASTE)
# ==========================================
with st.expander("📥 INPUT SOURCES (DOC / AUDIO STT / MIC / PASTE)", expanded=True):
    c_file, c_audio_stt, c_mic = st.columns([0.33, 0.34, 0.33])

    with c_file:
        st.markdown("**📁 DOC / TXT**")
        uploaded_file = st.file_uploader("Upload Doc", type=["docx", "txt"], key="doc_file_uploader", label_visibility="collapsed")
        if uploaded_file is not None:
            try:
                f_text = extract_text_from_file(uploaded_file)
                if f_text and f_text != st.session_state.main_text:
                    st.session_state.main_text = f_text
                    add_log(f"DOC Loaded: {uploaded_file.name}", "#4ade80")
                    st.toast(f"✅ {uploaded_file.name} Loaded!")
            except Exception as fe:
                st.error(f"Error: {fe}")

    with c_audio_stt:
        st.markdown("**🎵 AUDIO STT (Multi-min)**")
        stt_lang_choice = st.selectbox("Audio Input Lang:", options=["🔄 Auto (Multi-Lang)", "HI (हिंदी)", "TE (తెలుగు)", "EN (English)"], key="stt_lang_choice", label_visibility="collapsed")
        stt_lang_map = {"🔄 Auto (Multi-Lang)": "auto", "TE (తెలుగు)": "te-IN", "HI (हिंदी)": "hi-IN", "EN (English)": "en-IN"}
        selected_stt_lang = stt_lang_map[stt_lang_choice]

        uploaded_audio = st.file_uploader("Upload Audio", type=["mp3", "wav", "m4a", "ogg", "aac", "opus", "3gp"], key="audio_stt_file_uploader", label_visibility="collapsed")
        if uploaded_audio is not None:
            st.audio(uploaded_audio)
            use_dsp = st.checkbox("✨ DSP Booster", value=True, key="stt_dsp_chk")
            if st.button("🚀 RUN STT", use_container_width=True):
                add_log(f"STT Started: {uploaded_audio.name} ({stt_lang_choice})", "#c084fc")
                with st.spinner("Transcribing Voice..."):
                    raw_txt = transcribe_audio_file(uploaded_audio, lang_code=selected_stt_lang, enable_dsp=use_dsp)
                    if raw_txt and not raw_txt.startswith("⚠️"):
                        st.session_state.main_text = raw_txt.strip()
                        add_log(f"STT Ready ({len(raw_txt)} chars)", "#4ade80")
                        st.toast("✅ టెక్స్ట్ లోడ్ అయింది!")
                        st.rerun()
                    else:
                        st.error(raw_txt)

    with c_mic:
        st.markdown("**🎙️ LIVE MIC**")
        mic_lang = st.selectbox("Mic Lang:", options=["TE (తెలుగు)", "HI (हिंदी)", "EN (English)"], label_visibility="collapsed")
        mic_code_map = {"TE (తెలుగు)": "te-IN", "HI (हिंदी)": "hi-IN", "EN (English)": "en-IN"}
        spoken_result = speech_to_text(
            start_prompt="🎙️ START",
            stop_prompt="⏹️ STOP",
            language=mic_code_map[mic_lang],
            use_container_width=True,
            key='mic_rec'
        )
        if spoken_result and spoken_result != st.session_state.last_mic_text:
            st.session_state.main_text = (st.session_state.main_text + "\n\n" + spoken_result).strip()
            st.session_state.last_mic_text = spoken_result
            add_log(f"MIC: '{spoken_result}'", "#4ade80")
            st.rerun()


# ==========================================
# 4. బాక్స్ 1: మూల వచనం (Source Text Box)
# ==========================================
st.markdown("##### 📝 మూల వచనం (Source Text Box)")
user_input_text = st.text_area(
    "Original Source Content", 
    value=st.session_state.main_text, 
    height=120,
    placeholder="ఆడియో/మైక్/ఫైల్ నుంచి వచ్చిన లేదా ఇక్కడ నేరుగా పేస్ట్ చేసిన మూల వచనం...",
    label_visibility="collapsed"
)
if user_input_text != st.session_state.main_text:
    st.session_state.main_text = user_input_text


# ==========================================
# 5. రెండు బాక్సుల మధ్య ట్రాన్స్‌లేటర్ బార్ (Middle Translation Bar)
# ==========================================
col_mid_lbl, col_mid_lang, col_mid_btn = st.columns([0.25, 0.45, 0.3])

with col_mid_lbl:
    st.markdown("##### 🌐 ట్రాన్స్‌లేటర్ (AI):")

with col_mid_lang:
    target_trans_lang = st.selectbox(
        "మార్చాల్సిన భాష ఎంపిక:",
        options=[
            "హిందీ (सरल व आध्यात्मिक शैली)",
            "తెలుగు (గౌరవప్రదమైన కృష్ణా యాస / BK)",
            "ఇంగ్లీష్ (Dignified English)",
            "🔄 అసలు భాష (Original Polish Only)"
        ],
        index=0,
        label_visibility="collapsed"
    )

with col_mid_btn:
    if st.button("✨ అప్లై / అనువదించు (Apply Translation)", type="secondary", use_container_width=True):
        if st.session_state.main_text.strip():
            with st.spinner("ఎంచుకున్న భాషలోకి అనువదిస్తోంది..."):
                translated_res = polish_and_translate_script(
                    st.session_state.main_text, 
                    target_lang=target_trans_lang, 
                    style_mode=selected_style, 
                    pause_level=selected_pause
                )
                if translated_res:
                    st.session_state.translated_text = translated_res
                    add_log(f"అనువాదం పూర్తయింది -> {target_trans_lang}", "#38bdf8")
                    st.toast(f"✨ అనువాదం సిద్ధమైంది ({target_trans_lang})!", icon="✨")
                    st.rerun()
        else:
            st.warning("దయచేసి మూల వచనం బాక్స్‌లో టెక్స్ట్ ఎంటర్ చేయండి.")


# ==========================================
# 6. బాక్స్ 2: అనువాద ఫలితం (Final Translated Script)
# ==========================================
st.markdown("##### ✨ గౌరవప్రదమైన స్పీచ్ & అనువాద స్క్రిప్ట్ (Final Translated Script)")
trans_area = st.text_area(
    "Translated Content", 
    value=st.session_state.translated_text if st.session_state.translated_text else st.session_state.main_text, 
    height=130,
    label_visibility="collapsed"
)
if trans_area != st.session_state.translated_text:
    st.session_state.translated_text = trans_area

# ఆడియో & పోస్టర్ కోసం యాక్టివ్ టెక్స్ట్ నిర్ణయం
active_text = st.session_state.translated_text.strip() if st.session_state.translated_text.strip() else st.session_state.main_text.strip()


# ==========================================
# 7. TTS సెట్టింగ్స్
# ==========================================
with st.expander("⚙️ TTS SETTINGS (స్వరం, స్పీడ్ & BGM)", expanded=True):
    col_tts_lang, col_tts_voice = st.columns([0.45, 0.55])
    with col_tts_lang:
        tts_lang = st.selectbox("🌐 TTS Mode:", options=["🔄 Auto Detect (Multi-Lang)", "Hindi (हिंदी)", "Telugu (తెలుగు)", "English"], key="main_tts_lang_select")
    with col_tts_voice:
        gender_choice = st.radio("Voice Gender:", options=["👨 Male (పురుష)", "👩 Female (స్త్రీ)"], horizontal=True, key="gender_sel")

    col_opt_speed, col_opt_pitch, col_opt_pause = st.columns(3)
    with col_opt_speed:
        audio_speed = st.select_slider("🔊 Play Speed:", options=[0.75, 0.85, 1.0, 1.15, 1.25, 1.5], value=0.85, key="main_tts_speed")
    with col_opt_pitch:
        pitch_custom = st.select_slider("🎚️ Voice Pitch:", options=["Normal", "Deep Base", "Heavy Base"], value="Normal", key="main_tts_pitch")
    with col_opt_pause:
        pause_duration = st.slider("⏸️ Line Pause (Sec):", min_value=0.2, max_value=1.5, value=0.4, step=0.1, key="main_tts_pause")
        
    col_bgm_1, col_bgm_2 = st.columns([0.4, 0.6])
    with col_bgm_1:
        enable_bgm = st.checkbox("🎶 Enable BGM", value=True, key="main_tts_bgm_chk")
    with col_bgm_2:
        bgm_volume = st.slider("🎵 BGM Volume (%):", min_value=2, max_value=20, value=6, key="main_tts_bgm_vol")


# ==========================================
# 8. ప్రధాన యాక్షన్ కంట్రోల్స్
# ==========================================
b1, b2, b3, b4 = st.columns(4)
b5, b6, b7 = st.columns(3)

with b1:
    convert_btn = st.button("🔊 TTS (వాయిస్ చేయి)", type="primary", use_container_width=True)

with b2:
    if active_text:
        if st.button("🖼️ AI POSTER", use_container_width=True):
            with st.spinner("పోస్టర్ సిద్ధమవుతోంది..."):
                poster_html = generate_ai_poster_html(
                    active_text, 
                    theme=poster_theme, 
                    sticker_choice=sticker_choice, 
                    content_mode=content_mode,
                    text_align=text_align,
                    font_size_choice=font_size_choice,
                    custom_sticker_file=custom_sticker_file
                )
                st.session_state.poster_html_data = poster_html
                add_log("పోస్టర్ సిద్ధమైంది!", "#4ade80")
                st.toast("🖼️ పోస్టర్ సిద్ధమైంది!", icon="🖼️")
    else:
        st.button("🖼️ AI POSTER", disabled=True, use_container_width=True)

with b3:
    if active_text:
        html_trans_page = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body><p style='font-size:18px; line-height:1.8;'>{active_text.replace(chr(10), '<br>')}</p></body></html>"
        st.download_button("🌐 HTML", data=html_trans_page.encode('utf-8'), file_name="speech_script.html", mime="text/html", use_container_width=True)
    else:
        st.button("🌐 HTML", disabled=True, use_container_width=True)

with b4:
    if active_text:
        printable_pdf = create_printable_pdf_html(active_text)
        st.download_button("📄 PDF", data=printable_pdf.encode('utf-8'), file_name="speech_script.html", mime="text/html", use_container_width=True)
    else:
        st.button("📄 PDF", disabled=True, use_container_width=True)

with b5:
    if active_text:
        docx_data = create_docx_bytes(active_text)
        st.download_button("📝 DOCX", data=docx_data, file_name="speech_script.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    else:
        st.button("📝 DOCX", disabled=True, use_container_width=True)

with b6:
    if active_text:
        if st.button("📋 COPY", use_container_width=True):
            st.code(active_text, language=None)
            st.toast("✅ Copied!", icon="📋")
    else:
        st.button("📋 COPY", disabled=True, use_container_width=True)

with b7:
    if st.button("🧹 CLEAR", use_container_width=True):
        st.session_state.main_text = ""
        st.session_state.translated_text = ""
        st.session_state.audio_bytes_data = None
        st.session_state.poster_html_data = None
        st.session_state.last_mic_text = ""
        add_log("Cleared.", "#facc15")
        gc.collect()
        st.rerun()


# ==========================================
# 9. AI పోస్టర్ ప్రివ్యూ విభాగం
# ==========================================
if st.session_state.poster_html_data is not None:
    st.divider()
    st.markdown("### 🖼️ పోస్టర్ కార్డ్ ప్రివ్యూ (Smart Poster Card)")
    st.components.v1.html(st.session_state.poster_html_data, height=830, scrolling=True)


# ==========================================
# 10. TTS ఆడియో జనరేషన్
# ==========================================
if convert_btn:
    if active_text:
        add_log(f"TTS Started: Mode={tts_lang}", "#c084fc")
        with st.spinner("Generating Natural Voice..."):
            audio_bytes = synthesize_multilang_tts(
                text=active_text,
                tts_mode=tts_lang,
                gender=gender_choice,
                speed=audio_speed,
                pitch_custom=pitch_custom,
                pause_sec=pause_duration,
                enable_bgm=enable_bgm,
                bgm_vol=bgm_volume
            )
            if audio_bytes:
                st.session_state.audio_bytes_data = audio_bytes
                add_log("TTS Audio Ready!", "#4ade80")
                gc.collect()
                st.toast("🎉 TTS Audio Ready!")
            else:
                add_log("TTS Failed", "#f87171")
                st.error("❌ Audio Generation Failed.")
    else:
        st.warning("Please provide text.")

# ఆడియో ప్లేయర్ & డౌన్‌లోడ్
if st.session_state.audio_bytes_data is not None:
    st.divider()
    st.audio(st.session_state.audio_bytes_data, format="audio/mp3")
    st.download_button(
        label="📥 DOWNLOAD MP3", 
        data=st.session_state.audio_bytes_data, 
        file_name="speech_audio.mp3", 
        mime="audio/mp3",
        key="download_btn",
        use_container_width=True
    )

# డయాగ్నొస్టిక్స్
with st.expander("🔍 DIAGNOSTICS", expanded=False):
    log_html = "<div class='diag-box'>"
    for item in st.session_state.diag_logs[-15:]:
        log_html += f"<div class='diag-log'><span style='color:#94a3b8;'>[{item['time']}]</span> <span style='color:{item['color']};'>{item['msg']}</span></div>"
    log_html += "</div>"
    st.markdown(log_html, unsafe_allow_html=True)
